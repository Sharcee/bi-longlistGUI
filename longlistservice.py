###################
#   Imports       #
###################
import sys

from datetime import datetime
from docx.shared import Pt
from docx.shared import Inches
from docx import Document
import os

from openpyxl import load_workbook

###################
#  Main Routines  #
###################
def fillHeaderInfo(doc, data, section):
    _section = section
    header = _section.header
    headers = header.paragraphs
    paragraphs = list(headers)

    for p in paragraphs:
        for key, val in data.items():
            key_name = key
            if key_name in p.text:
                inline = p.runs
                # Replace strings and retain the same style.
                # The text to be replaced can be split over several runs so
                # search through, identify which runs need to have text replaced
                # then replace the text in those identified
                started = False
                key_index = 0
                # found_runs is a list of (inline index, index of match, length of match)
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(inline)):

                    # case 1: found in single run so short circuit the replace
                    if key_name in inline[i].text and not started:
                        found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        text = inline[i].text.replace(key_name, str(val))
                        inline[i].text = text
                        replace_done = True
                        found_all = True
                        break

                    if key_name[key_index] not in inline[i].text and not started:
                        # keep looking ...
                        continue

                    # case 2: search for partial text, find first run
                    if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                        # check sequence
                        start_index = inline[i].text.find(key_name[key_index])
                        check_length = len(inline[i].text)
                        for text_index in range(start_index, check_length):
                            if inline[i].text[text_index] != key_name[key_index]:
                                # no match so must be false positive
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key_name):
                            continue
                        else:
                            # found all chars in key_name
                            found_all = True
                            break

                    # case 2: search for partial text, find subsequent run
                    if key_name[key_index] in inline[i].text and started and not found_all:
                        # check sequence
                        chars_found = 0
                        check_length = len(inline[i].text)
                        for text_index in range(0, check_length):
                            if inline[i].text[text_index] == key_name[key_index]:
                                key_index += 1
                                chars_found += 1
                            else:
                                break
                        # no match so must be end
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key_name):
                            found_all = True
                            break

                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            inline[index].text = text
                        else:
                            text = inline[index].text.replace(inline[index].text[start:start + length], '')
                            inline[index].text = text
                # print(p.text)

def addFormalTitle(doc, company, project, date):
    title = doc.paragraphs[0]
    title.clear()

    # Set Paragraph Text
    title_run = title.add_run(company.upper() + " PROJECT " + '"' + project.upper() + '"\nNAME CANDIDATE LONGLIST - ' + date.upper())
    # Set Font, Size
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    # Single Spaced; Pt(6) font spacing after
    title.paragraph_format.line_spacing = 1
    title.paragraph_format.space_after = Pt(6)

def addCategory(doc, data):
    addSizeTenBuffer(doc)
    p = initParagraph(doc)
    pr = p.add_run(data)
    pr.font.name = 'Calibri'
    pr.font.size = Pt(14)
    pr.underline = True
    pr.bold = True

def addBulletName(doc, name, rationale, symbol="", team=False):
    TAB = '\t'
    bullet = doc.add_paragraph(style='mynew')
    bullrun = bullet.add_run(name)
    subrun = bullet.add_run(symbol + TAB)

    if team:
        bullrun.bold = True
        subrun.bold = True
    
    bullrun.font.size = Pt(14)
    bullrun.font.name = 'Calibri'
    subrun.font.size = Pt(8)

    ratrun = bullet.add_run(rationale)
    ratrun.font.size = Pt(10)
    ratrun.font.name = 'Calibri'


# def addScreenerInfo(doc, data, section):
#     pass

###################
# Routine Helpers #
###################
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def convertDate(mmddyy):
    dateobj = datetime.strptime(mmddyy, '%m.%d.%y')
    return dateobj.strftime('%B %d, %Y').replace(" 0", " ")

def buildClientDict(company, project, date):
    tmp = 'Project ' + project + ' Longlist'
    return {
        'Project Project Name Longlist' : tmp,
        'Company' : company,
        'Date' : date
    }

def initParagraph(doc):
    tmp = doc.add_paragraph()
    tmp.paragraph_format.space_before = None
    tmp.paragraph_format.space_after = Pt(8)

    return tmp

def addSizeTenBuffer(doc):
    x = doc.add_paragraph()
    x_run = x.add_run(" ")
    x_run.font.name = 'Calibri'
    x_run.font.size = Pt(10)
    x.paragraph_format.space_after = Pt(6)
    pass

###################
#      Main       #
###################
def main(filePath):
    # Working with input file name
    folder = filePath.split("/")
    myFile = folder[-1]
    folder = "/".join(folder[:-1])+"/"

    company, project, _, date = myFile.split("_")
    date = '.'.join(date.split('.')[:-1])

    # Create document object from template
    print(resource_path('ll_template'))
    doc = Document(resource_path('ll_template.docx'))

    # Replace our Header
    section = doc.sections[0]
    fillHeaderInfo(doc, buildClientDict(company, project.upper(), convertDate(date)), section)

    # Place our Title
    addFormalTitle(doc, company, project, convertDate(date))

    # Load Workbook
    wb = load_workbook(filename = filePath)
    ws = wb.active

    main_category = None
    main_category_name = None
    sub_category = None
    containsSubcategories = True

    for row in ws.iter_rows(min_row=2):
        seq, category, name, rationale = row
        
        if str(seq.value).isalpha():
            main_category = seq.value
            main_category_name = category.value
            sub_category = None

            addCategory(doc, main_category+'. '+main_category_name)
            continue
        
        if not category.value:
            containsSubcategories = False

        if(containsSubcategories):
            if sub_category is None or category.value != sub_category:
                sub_category = category.value
                addCategory(doc, sub_category)

        sub = ""
        if "(" in name.value:
            name.value, sub = name.value.split("(", 1)
            sub = "(" + sub

        addBulletName(doc, name.value, rationale.value, sub, name.font.bold)
    
    for _par in doc.paragraphs:
        for q in _par.runs:
            if "'" in q.text:
                q.text = q.text.replace("'", "’")
    
    # Save into our word doc
    filePath = "/".join(filePath.split("/")[:-1])+"/"
    # print(filePath)
    our_file = filePath + "{} {} Longlist {}.docx".format(company,project.upper(), date)
    
    doc.save(our_file)
    # print(our_file)

    return our_file

# if __name__ == "__main__":
#     pass

