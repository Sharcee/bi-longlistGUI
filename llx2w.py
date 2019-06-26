###################
#   Imports       #
###################
import sys

from datetime import datetime
from docx.shared import Pt
from docx.shared import Inches
from docx import Document
import os

from openpyxl import load_workbook

def addCategoryTitle(doc, letter, name):
    text = "{}. {}".format(letter, name)
    size10buffer(doc)
    # Define our Category Title Formatting
    category_section_title = doc.add_paragraph()
    category_section_title.paragraph_format.space_before = None
    category_section_title.paragraph_format.space_after = Pt(8)
    # Title
    category = category_section_title.add_run(text)
    category.font.name = 'Calibri'
    category.font.size = Pt(14)
    category.underline = True
    category.bold = True

def addSubTitle(doc, name):
    size10buffer(doc)
    # Define our Category Title Formatting
    category_section_title = doc.add_paragraph()
    category_section_title.paragraph_format.space_before = None
    category_section_title.paragraph_format.space_after = Pt(8)
    # Title
    category = category_section_title.add_run(name)
    category.font.name = 'Calibri'
    category.font.size = Pt(14)
    category.underline = True
    category.bold = True

def size10buffer(doc):
    title_space = doc.add_paragraph()
    title_space_run = title_space.add_run(" ")
    title_space_run.font.name = 'Calibri'
    title_space_run.font.size = Pt(10)
    title_space.paragraph_format.space_after = Pt(6)

def convertDate(mmddyy):
    dateobj = datetime.strptime(mmddyy, '%m.%d.%y')
    return dateobj.strftime('%B %d, %Y').replace(" 0", " ")

def buildReplaceDict(company, project, date):
    temp = 'Project ' + project + ' Longlist'
    return {
        'Project Project Name Longlist' : temp,
        'Company' : company,
        'Date' : date
    }


#@TODO: Refactor this Code you found on StackOverflow
def docx_replace_header(doc, data, section):
    _section = section
    header = _section.header
    headers = header.paragraphs
    paragraphs = list(headers)

    for p in paragraphs:
        for key, val in data.items():
            key_name = key
            if key_name in p.text:
                inline = p.runs
                # Replace strings and retain the same style.
                # The text to be replaced can be split over several runs so
                # search through, identify which runs need to have text replaced
                # then replace the text in those identified
                started = False
                key_index = 0
                # found_runs is a list of (inline index, index of match, length of match)
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(inline)):

                    # case 1: found in single run so short circuit the replace
                    if key_name in inline[i].text and not started:
                        found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        text = inline[i].text.replace(key_name, str(val))
                        inline[i].text = text
                        replace_done = True
                        found_all = True
                        break

                    if key_name[key_index] not in inline[i].text and not started:
                        # keep looking ...
                        continue

                    # case 2: search for partial text, find first run
                    if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                        # check sequence
                        start_index = inline[i].text.find(key_name[key_index])
                        check_length = len(inline[i].text)
                        for text_index in range(start_index, check_length):
                            if inline[i].text[text_index] != key_name[key_index]:
                                # no match so must be false positive
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key_name):
                            continue
                        else:
                            # found all chars in key_name
                            found_all = True
                            break

                    # case 2: search for partial text, find subsequent run
                    if key_name[key_index] in inline[i].text and started and not found_all:
                        # check sequence
                        chars_found = 0
                        check_length = len(inline[i].text)
                        for text_index in range(0, check_length):
                            if inline[i].text[text_index] == key_name[key_index]:
                                key_index += 1
                                chars_found += 1
                            else:
                                break
                        # no match so must be end
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key_name):
                            found_all = True
                            break

                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            inline[index].text = text
                        else:
                            text = inline[index].text.replace(inline[index].text[start:start + length], '')
                            inline[index].text = text
                # print(p.text)




def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def main(fileIN):
    """
    Input Template Example: company_PROJECT_N 4 LL for upload_MM.DD.YY.xlsx
    print(fileIN.split('_'))
    ['company', 'PROJECT', 'N4LL for upload', 'MM.DD.YY.xlsx']

    Header Section
    """
    relative = fileIN.split("/")[-1]
    # Store Header Names
    company, project, _, date = relative.split("_")

    # Clean Up Date: 05.05.95.docx -> 05.05.95
    date = '.'.join(date.split(".")[:-1])
    date1 = convertDate(date)
    # print(company, project, date)

    # Opening our document
    doc = Document(resource_path('ll_template.docx'))

    # Edit our Headers
    section = doc.sections[0]
    docx_replace_header(doc, buildReplaceDict(company, project.upper(), date1), section)

    """
    LongList Title Section
    """
    # Select First Paragraph from Template
    title = doc.paragraphs[0]
    title.clear()

    # Set Paragraph Text
    title_run = title.add_run(company.upper() + " PROJECT " + '"' + project.upper() + '"\nNAME CANDIDATE LONGLIST - ' + date1.upper())
    # Set Font, Size
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    # Single Spaced; Pt(6) font spacing after
    title.paragraph_format.line_spacing = 1
    title.paragraph_format.space_after = Pt(6)


    """
    Scrape Excel Sheet for Data
    """
    # Load Workbook
    wb = load_workbook(filename = fileIN)
    ws = wb.active

    main_category = None
    main_category_name = None
    sub_category = None

    for row in ws.iter_rows(min_row = 2):
        seq, category, name, rationale = row

        if str(seq.value).isalpha():
            main_category = seq.value
            main_category_name = category.value
            sub_category = None

            addCategoryTitle(doc, main_category, main_category_name)
            continue
        
        if sub_category is None or category.value != sub_category:
            sub_category = category.value
            addSubTitle(doc, sub_category)
        
        bullet = doc.add_paragraph(style='mynew')
        TABS = "\t"
        sub = ""
        BOLD_INDICATOR = '$'

        if name.font.bold:
            name = BOLD_INDICATOR + name.value.strip()
        else:
            name = name.value            

        if "(" in name:
            name, sub = name.split("(")
            sub = "(" + sub

        if BOLD_INDICATOR in name:
            namerun = bullet.add_run(name[1:])  # UGLY -- make pythonic
            subrun = bullet.add_run(sub + TABS)
            namerun.bold = True # Bold our Run
            subrun.bold = True
        else:
            namerun = bullet.add_run(name)
            subrun = bullet.add_run(sub + TABS)
        
        subrun.font.size = Pt(8)
        namerun.font.size = Pt(14)
        namerun.font.name = 'Calibri'

        # Apply Font
        rational = bullet.add_run(rationale.value)
        rational.font.size = Pt(10)
        rational.font.name = 'Calibri'

    for _par in doc.paragraphs:
        for q in _par.runs:
            if "'" in q.text:
                q.text = q.text.replace("'", "’")
            # if '"' in q.text:
            #     q.text = q.text.replace('"', '“')


    # Save into our word doc
    our_file = "{} {} Longlist {}.docx".format(company,project.upper(), date)
    doc.save(our_file)

    return our_file

# if __name__ == "__main__":
#     main(sys.argv[1])